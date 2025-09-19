"""
DOC/DOCX文字提取服务模块

基于python-docx和python-docx2txt实现DOC/DOCX文字提取功能，支持：
1. 原生文本提取
2. 图像检测和提取
3. 图文混合文档处理
4. OCR和Qwen-VL双线处理
5. 结构化输出
"""

import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import asyncio
import concurrent.futures
from datetime import datetime
import uuid
import io
from PIL import Image

# DOC/DOCX处理库
from docx import Document
import docx2txt
import subprocess
import tempfile
import os

# 现有服务
from .tesseract_ocr_service import tesseract_ocr_service
from .baidu_ocr_service import baidu_ocr_service
from ..schemas.ocr_schemas import OCREngineType
from .vision_service import vision_service
from ..schemas.doc_schemas import (
    DocType, DocExtractionMethod, DocExtractionResult, 
    DocPageInfo, DocParagraphInfo, DocImageInfo, 
    DocProcessingStats, DocContentAnalysis, DocImageExtractionResult
)

logger = logging.getLogger(__name__)


class DocTextExtractor:
    """DOC/DOCX文字提取器"""
    
    def __init__(self):
        """初始化DOC提取器"""
        self.min_text_length = 10  # 最小文本长度阈值
        
    def extract_text_from_doc(self, doc_path: str, doc_type: DocType, 
                             use_ocr: bool = False, ocr_engine: OCREngineType = OCREngineType.BAIDU,
                             use_vision: bool = False, vision_model: str = "qwen-vl-plus",
) -> DocExtractionResult:
        """
        从DOC/DOCX文件中提取文字 - 实现四步法处理流程
        
        步骤1: 检测文档类型和内容
        步骤2: 提取原生文字和图像
        步骤3: 对图像进行OCR和视觉模型处理
        步骤4: 融合结果并输出
        
        Args:
            doc_path: DOC/DOCX文件路径
            doc_type: 文档类型 (DOC/DOCX)
            use_ocr: 是否对图像使用OCR
            ocr_engine: OCR引擎类型
            use_vision: 是否使用Qwen-VL多模态模型
            vision_model: 视觉模型名称
            
        Returns:
            包含提取结果的DocExtractionResult对象
        """
        start_time = datetime.now()
        
        try:
            logger.info(f"开始处理{doc_type.value.upper()}文档: {doc_path}")
            
            # 步骤1: 检测文档类型和内容
            content_analysis = self._analyze_doc_content(doc_path, doc_type)
            logger.info(f"文档内容分析结果: {content_analysis}")
            
            # 步骤2: 提取原生文字和图像
            if doc_type == DocType.DOCX:
                extraction_result = self._extract_from_docx(
                    doc_path, content_analysis, use_ocr, ocr_engine, 
                    use_vision, vision_model, start_time
                )
            else:  # DOC
                extraction_result = self._extract_from_doc(
                    doc_path, content_analysis, use_ocr, ocr_engine,
                    use_vision, vision_model, start_time
                )
            
            # 计算处理时间
            processing_time = (datetime.now() - start_time).total_seconds()
            extraction_result.processing_time = processing_time
            
            logger.info(f"DOC文档处理完成，耗时: {processing_time:.2f}秒")
            return extraction_result
            
        except Exception as e:
            logger.error(f"DOC文档文字提取失败: {doc_path}, 错误: {str(e)}")
            raise Exception(f"DOC文档文字提取失败: {str(e)}")
    
    def _analyze_doc_content(self, doc_path: str, doc_type: DocType) -> DocContentAnalysis:
        """
        分析DOC文档内容
        
        Args:
            doc_path: 文档路径
            doc_type: 文档类型
            
        Returns:
            内容分析结果
        """
        try:
            has_images = False
            has_native_text = False
            total_images = 0
            total_text_length = 0
            
            if doc_type == DocType.DOCX:
                # 分析DOCX文档
                doc = Document(doc_path)
                
                # 检查文本内容
                full_text = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        full_text += paragraph.text + "\n"
                
                total_text_length = len(full_text.strip())
                has_native_text = total_text_length > self.min_text_length
                
                # 检查图像
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        has_images = True
                        total_images += 1
                        
            else:  # DOC
                # 使用docx2txt分析DOC文档
                try:
                    text_content = docx2txt.process(doc_path)
                    total_text_length = len(text_content.strip())
                    has_native_text = total_text_length > self.min_text_length
                except Exception as e:
                    logger.warning(f"DOC文档文本提取失败: {str(e)}")
                    has_native_text = False
                
                # DOC格式的图像检测比较复杂，这里简化处理
                has_images = True  # 保守估计
                total_images = 1
            
            # 判断处理策略
            use_ocr = not has_native_text or has_images
            use_vision = has_images
            include_images = has_images
            
            return DocContentAnalysis(
                has_images=has_images,
                has_native_text=has_native_text,
                total_pages=1,  # DOC文档通常不分页
                pages_with_images=1 if has_images else 0,
                pages_with_text=1 if has_native_text else 0,
                total_images=total_images,
                use_ocr=use_ocr,
                use_vision=use_vision,
                include_images=include_images,
                detection_confidence="high" if doc_type == DocType.DOCX else "medium"
            )
            
        except Exception as e:
            logger.error(f"DOC内容分析失败: {str(e)}")
            return DocContentAnalysis(
                has_images=True,
                has_native_text=False,
                total_pages=1,
                pages_with_images=0,
                pages_with_text=0,
                total_images=0,
                use_ocr=True,
                use_vision=True,
                include_images=True,
                detection_confidence="low",
                error=str(e)
            )
    
    def _extract_from_docx(self, doc_path: str, content_analysis: DocContentAnalysis,
                          use_ocr: bool, ocr_engine: OCREngineType,
                          use_vision: bool, vision_model: str,
                          start_time: datetime,
) -> DocExtractionResult:
        """
        从DOCX文档提取内容
        
        Args:
            doc_path: DOCX文件路径
            content_analysis: 内容分析结果
            use_ocr: 是否使用OCR
            ocr_engine: OCR引擎
            use_vision: 是否使用视觉模型
            vision_model: 视觉模型名称
            
        Returns:
            提取结果
        """
        try:
            doc = Document(doc_path)
            
            # 提取文档属性
            doc_properties = {
                "title": doc.core_properties.title,
                "author": doc.core_properties.author,
                "subject": doc.core_properties.subject,
                "created": doc.core_properties.created.isoformat() if doc.core_properties.created else None,
                "modified": doc.core_properties.modified.isoformat() if doc.core_properties.modified else None,
                "last_modified_by": doc.core_properties.last_modified_by,
                "revision": doc.core_properties.revision,
                "version": doc.core_properties.version
            }
            
            # 提取段落和图像
            paragraphs = []
            images = []
            full_text = ""
            
            # 处理段落
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    para_info = self._extract_paragraph_info(paragraph, para_idx)
                    paragraphs.append(para_info)
                    full_text += paragraph.text + "\n"
            
            # 提取图像并进行OCR/视觉处理
            images = []
            image_texts = []  # 存储图像识别结果
            
            if content_analysis.has_images:
                images, image_texts = self._extract_images_from_docx_with_fusion(
                    doc, use_ocr, ocr_engine, use_vision, vision_model
                )
            
            # 合并原生文本和图像识别文本
            combined_text = full_text.strip()
            if image_texts:
                combined_text += "\n\n[图像文字识别结果]\n" + "\n".join(image_texts)
            
            # 创建页面信息（DOCX通常作为单页处理）
            page_info = DocPageInfo(
                page_number=1,
                text=combined_text,
                text_length=len(combined_text.strip()),
                paragraphs=paragraphs,
                images=images,
                processing_type=DocExtractionMethod.MIXED_CONTENT if images else DocExtractionMethod.NATIVE_TEXT,
                extraction_method="docx_native_extraction_with_fusion" if image_texts else "docx_native_extraction",
                has_images=len(images) > 0
            )
            
            # 统计信息
            stats = DocProcessingStats(
                total_pages=1,
                native_text_pages=1 if content_analysis.has_native_text else 0,
                ocr_processed_pages=1 if use_ocr and images else 0,
                vision_processed_pages=1 if use_vision and images else 0,
                image_only_pages=0,
                mixed_content_pages=1 if images else 0,
                error_pages=0,
                total_images=len(images),
                processed_images=len([img for img in images if img.has_ocr_text or img.has_vision_text])
            )
            
            return DocExtractionResult(
                file_path=doc_path,
                file_type=DocType.DOCX,
                total_pages=1,
                pages=[page_info],
                full_text=combined_text,
                has_text_layer=content_analysis.has_native_text,
                is_scanned=False,  # DOCX不是扫描文档
                extraction_method="docx_four_step_process",
                processing_stats=stats,
                document_properties=doc_properties,
                created_at=start_time.isoformat()
            )
            
        except Exception as e:
            logger.error(f"DOCX文档提取失败: {str(e)}")
            raise
    
    def _extract_from_doc(self, doc_path: str, content_analysis: DocContentAnalysis,
                         use_ocr: bool, ocr_engine: OCREngineType,
                         use_vision: bool, vision_model: str,
                         start_time: datetime,
) -> DocExtractionResult:
        """
        从DOC文档提取内容
        
        Args:
            doc_path: DOC文件路径
            content_analysis: 内容分析结果
            use_ocr: 是否使用OCR
            ocr_engine: OCR引擎
            use_vision: 是否使用视觉模型
            vision_model: 视觉模型名称
            
        Returns:
            提取结果
        """
        try:
            # 使用docx2txt提取文本
            text_content = docx2txt.process(doc_path)
            
            # 创建段落信息
            paragraphs = []
            if text_content.strip():
                para_info = DocParagraphInfo(
                    paragraph_id="para_1",
                    text=text_content.strip(),
                    text_length=len(text_content.strip()),
                    font_name=None,
                    font_size=None,
                    is_bold=False,
                    is_italic=False,
                    is_underlined=False,
                    alignment=None,
                    has_images=False,
                    image_ids=[]
                )
                paragraphs.append(para_info)
            
            # 对于DOC格式，图像提取比较复杂，这里简化处理
            images = []
            if content_analysis.has_images and (use_ocr or use_vision):
                # 尝试提取图像（这里需要更复杂的实现）
                logger.info("DOC格式图像提取功能待完善")
            
            # 创建页面信息
            page_info = DocPageInfo(
                page_number=1,
                text=text_content.strip(),
                text_length=len(text_content.strip()),
                paragraphs=paragraphs,
                images=images,
                processing_type=DocExtractionMethod.NATIVE_TEXT,
                extraction_method="doc_native_extraction",
                has_images=len(images) > 0
            )
            
            # 统计信息
            stats = DocProcessingStats(
                total_pages=1,
                native_text_pages=1 if content_analysis.has_native_text else 0,
                ocr_processed_pages=0,
                vision_processed_pages=0,
                image_only_pages=0,
                mixed_content_pages=1 if images else 0,
                error_pages=0,
                total_images=len(images),
                processed_images=0
            )
            
            return DocExtractionResult(
                file_path=doc_path,
                file_type=DocType.DOC,
                total_pages=1,
                pages=[page_info],
                full_text=text_content.strip(),
                has_text_layer=content_analysis.has_native_text,
                is_scanned=False,
                extraction_method="doc_four_step_process",
                processing_stats=stats,
                document_properties=None,
                created_at=start_time.isoformat()
            )
            
        except Exception as e:
            logger.error(f"DOC文档提取失败: {str(e)}")
            raise
    
    def _extract_paragraph_info(self, paragraph, para_idx: int) -> DocParagraphInfo:
        """
        提取段落信息
        
        Args:
            paragraph: docx段落对象
            para_idx: 段落索引
            
        Returns:
            段落信息
        """
        try:
            para_info = DocParagraphInfo(
                paragraph_id=f"para_{para_idx}",
                text=paragraph.text,
                text_length=len(paragraph.text),
                has_images=False,
                image_ids=[]
            )
            
            
            return para_info
            
        except Exception as e:
            logger.error(f"段落信息提取失败: {str(e)}")
            return DocParagraphInfo(
                paragraph_id=f"para_{para_idx}",
                text=paragraph.text,
                text_length=len(paragraph.text),
                has_images=False,
                image_ids=[]
            )
    
    def _extract_images_from_docx(self, doc: Document, use_ocr: bool, ocr_engine: OCREngineType,
                                 use_vision: bool, vision_model: str) -> List[DocImageInfo]:
        """
        从DOCX文档提取图像并进行OCR/视觉处理 - 实现与PDF一致的双线并行处理
        
        Args:
            doc: docx文档对象
            use_ocr: 是否使用OCR
            ocr_engine: OCR引擎
            use_vision: 是否使用视觉模型
            vision_model: 视觉模型名称
            
        Returns:
            图像信息列表
        """
        images = []
        
        try:
            # 获取所有关系中的图像
            for rel_idx, rel in enumerate(doc.part.rels.values()):
                if "image" in rel.target_ref:
                    try:
                        # 提取图像数据
                        image_data = rel.target_part.blob
                        image_id = f"img_{rel_idx}"
                        
                        # 创建图像信息
                        img_info = DocImageInfo(
                            image_id=image_id,
                            image_name=f"image_{rel_idx}",
                            image_type=rel.target_ref.split('.')[-1].lower(),
                            image_size=len(image_data),
                            has_ocr_text=False,
                            has_vision_text=False
                        )
                        
                        # 如果启用OCR或视觉处理，进行图像识别
                        if use_ocr or use_vision:
                            logger.info(f"开始处理图像{image_id}: use_ocr={use_ocr}, use_vision={use_vision}")
                            
                            # 使用线程池进行真正的并行处理
                            with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
                                futures = {}
                                
                                # 启动OCR处理任务
                                if use_ocr:
                                    logger.info(f"图像{image_id}启动OCR处理任务")
                                    ocr_future = executor.submit(
                                        self._process_image_with_ocr, 
                                        image_data, ocr_engine
                                    )
                                    futures['ocr'] = ocr_future
                                else:
                                    logger.info(f"图像{image_id}OCR未启用")
                                
                                # 启动视觉模型处理任务
                                if use_vision and vision_service.is_available():
                                    logger.info(f"图像{image_id}启动Qwen-VL处理任务")
                                    vision_future = executor.submit(
                                        self._process_image_with_vision, 
                                        image_data, vision_model
                                    )
                                    futures['vision'] = vision_future
                                else:
                                    if not use_vision:
                                        logger.info(f"图像{image_id}Qwen-VL未启用")
                                    else:
                                        logger.warning(f"图像{image_id}Qwen-VL服务不可用")
                                
                                # 等待所有任务完成
                                ocr_result = None
                                vision_result = None
                                
                                for task_name, future in futures.items():
                                    try:
                                        if task_name == 'ocr':
                                            ocr_result = future.result(timeout=30)  # 30秒超时
                                            logger.info(f"图像{image_id}OCR处理完成: {ocr_result.get('success', False)}")
                                        elif task_name == 'vision':
                                            vision_result = future.result(timeout=60)  # 60秒超时
                                            logger.info(f"图像{image_id}Qwen-VL处理完成: {vision_result.get('success', False)}")
                                    except concurrent.futures.TimeoutError:
                                        logger.warning(f"图像{image_id}的{task_name}处理超时")
                                    except Exception as e:
                                        logger.error(f"图像{image_id}的{task_name}处理异常: {str(e)}")
                                
                                # 融合决策模块：择优/纠错/补全
                                fused_text, extraction_method = self._fusion_decision_module_for_image(
                                    ocr_result, vision_result, image_id
                                )
                                
                                # 更新图像信息
                                if ocr_result and ocr_result.get('success'):
                                    img_info.has_ocr_text = True
                                if vision_result and vision_result.get('success'):
                                    img_info.has_vision_text = True
                                
                                # 存储融合结果（如果需要的话）
                                if fused_text:
                                    logger.info(f"图像{image_id}融合结果: {extraction_method}")
                        
                        images.append(img_info)
                        
                    except Exception as e:
                        logger.error(f"图像{rel_idx}提取失败: {str(e)}")
                        continue
            
            logger.info(f"成功提取{len(images)}个图像")
            return images
            
        except Exception as e:
            logger.error(f"DOCX图像提取失败: {str(e)}")
            return []
    
    def _extract_images_from_docx_with_fusion(self, doc: Document, use_ocr: bool, ocr_engine: OCREngineType,
                                            use_vision: bool, vision_model: str) -> Tuple[List[DocImageInfo], List[str]]:
        """
        从DOCX文档提取图像并进行OCR/视觉处理，返回融合后的文本结果
        
        Args:
            doc: docx文档对象
            use_ocr: 是否使用OCR
            ocr_engine: OCR引擎
            use_vision: 是否使用视觉模型
            vision_model: 视觉模型名称
            
        Returns:
            (图像信息列表, 融合后的文本列表)
        """
        images = []
        image_texts = []
        
        try:
            # 获取所有关系中的图像
            for rel_idx, rel in enumerate(doc.part.rels.values()):
                if "image" in rel.target_ref:
                    try:
                        # 提取图像数据
                        image_data = rel.target_part.blob
                        image_id = f"img_{rel_idx}"
                        
                        # 创建图像信息
                        img_info = DocImageInfo(
                            image_id=image_id,
                            image_name=f"image_{rel_idx}",
                            image_type=rel.target_ref.split('.')[-1].lower(),
                            image_size=len(image_data),
                            has_ocr_text=False,
                            has_vision_text=False
                        )
                        
                        # 如果启用OCR或视觉处理，进行图像识别
                        if use_ocr or use_vision:
                            logger.info(f"开始处理图像{image_id}: use_ocr={use_ocr}, use_vision={use_vision}")
                            
                            # 使用线程池进行真正的并行处理
                            with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
                                futures = {}
                                
                                # 启动OCR处理任务
                                if use_ocr:
                                    logger.info(f"图像{image_id}启动OCR处理任务")
                                    ocr_future = executor.submit(
                                        self._process_image_with_ocr, 
                                        image_data, ocr_engine
                                    )
                                    futures['ocr'] = ocr_future
                                else:
                                    logger.info(f"图像{image_id}OCR未启用")
                                
                                # 启动视觉模型处理任务
                                if use_vision and vision_service.is_available():
                                    logger.info(f"图像{image_id}启动Qwen-VL处理任务")
                                    vision_future = executor.submit(
                                        self._process_image_with_vision, 
                                        image_data, vision_model
                                    )
                                    futures['vision'] = vision_future
                                else:
                                    if not use_vision:
                                        logger.info(f"图像{image_id}Qwen-VL未启用")
                                    else:
                                        logger.warning(f"图像{image_id}Qwen-VL服务不可用")
                                
                                # 等待所有任务完成
                                ocr_result = None
                                vision_result = None
                                
                                for task_name, future in futures.items():
                                    try:
                                        if task_name == 'ocr':
                                            ocr_result = future.result(timeout=30)  # 30秒超时
                                            logger.info(f"图像{image_id}OCR处理完成: {ocr_result.get('success', False)}")
                                        elif task_name == 'vision':
                                            vision_result = future.result(timeout=60)  # 60秒超时
                                            logger.info(f"图像{image_id}Qwen-VL处理完成: {vision_result.get('success', False)}")
                                    except concurrent.futures.TimeoutError:
                                        logger.warning(f"图像{image_id}的{task_name}处理超时")
                                    except Exception as e:
                                        logger.error(f"图像{image_id}的{task_name}处理异常: {str(e)}")
                                
                                # 融合决策模块：择优/纠错/补全
                                fused_text, extraction_method = self._fusion_decision_module_for_image(
                                    ocr_result, vision_result, image_id
                                )
                                
                                # 更新图像信息
                                if ocr_result and ocr_result.get('success'):
                                    img_info.has_ocr_text = True
                                if vision_result and vision_result.get('success'):
                                    img_info.has_vision_text = True
                                
                                # 存储融合结果（用户友好格式）
                                if fused_text and fused_text.strip():
                                    clean_text = self._format_user_friendly_text(fused_text, extraction_method)
                                    if clean_text:  # 只有有意义的文本才添加
                                        image_texts.append(clean_text)
                                    logger.info(f"图像{image_id}融合结果: {extraction_method}")
                        
                        images.append(img_info)
                        
                    except Exception as e:
                        logger.error(f"图像{rel_idx}提取失败: {str(e)}")
                        continue
            
            logger.info(f"成功提取{len(images)}个图像，生成{len(image_texts)}个文本结果")
            return images, image_texts
            
        except Exception as e:
            logger.error(f"DOCX图像提取失败: {str(e)}")
            return [], []
    
    def _process_image_with_ocr(self, image_data: bytes, ocr_engine: OCREngineType) -> Dict[str, Any]:
        """
        使用OCR处理图像
        
        Args:
            image_data: 图像数据
            ocr_engine: OCR引擎
            
        Returns:
            OCR处理结果
        """
        try:
            # 转换为PIL图像
            img_pil = Image.open(io.BytesIO(image_data))
            
            # 进行OCR识别
            if ocr_engine == OCREngineType.TESSERACT:
                result = tesseract_ocr_service.extract_text(img_pil)
            elif ocr_engine == OCREngineType.BAIDU:
                result = baidu_ocr_service.extract_text(img_pil)
            else:
                raise ValueError(f"不支持的OCR引擎: {ocr_engine}")
            
            return result
            
        except Exception as e:
            logger.error(f"图像OCR识别失败: {str(e)}")
            return {"text": "", "error": str(e), "success": False}
    
    def _process_image_with_vision(self, image_data: bytes, vision_model: str) -> Dict[str, Any]:
        """
        使用视觉模型处理图像
        
        Args:
            image_data: 图像数据
            vision_model: 视觉模型名称
            
        Returns:
            视觉模型处理结果
        """
        try:
            # 使用Qwen-VL进行识别
            result = vision_service.extract_text_from_image(image_data, vision_model)
            return result
            
        except Exception as e:
            logger.error(f"图像视觉模型识别失败: {str(e)}")
            return {"text": "", "error": str(e), "success": False}
    
    def _fusion_decision_module_for_image(self, ocr_result: Optional[Dict], vision_result: Optional[Dict], 
                                        image_id: str) -> Tuple[str, str]:
        """
        融合决策模块：择优/纠错/补全 - 针对单个图像
        
        Args:
            ocr_result: OCR处理结果
            vision_result: Qwen-VL处理结果
            image_id: 图像ID
            
        Returns:
            (融合文本, 提取方法)
        """
        try:
            # 获取结果状态
            ocr_success = ocr_result and ocr_result.get("success", False)
            vision_success = vision_result and vision_result.get("success", False)
            
            ocr_text = ocr_result.get("text", "").strip() if ocr_result else ""
            vision_text = vision_result.get("text", "").strip() if vision_result else ""
            
            ocr_confidence = ocr_result.get("confidence", 0) if ocr_result else 0
            vision_confidence = vision_result.get("confidence", 0) if vision_result else 0
            
            logger.info(f"图像{image_id}融合决策: OCR成功={ocr_success}, Vision成功={vision_success}")
            logger.info(f"图像{image_id}文本长度: OCR={len(ocr_text)}, Vision={len(vision_text)}")
            
            # 决策逻辑
            if ocr_success and vision_success:
                # 两者都成功：智能融合
                logger.info(f"图像{image_id}两者都成功，进行智能融合")
                return self._intelligent_fusion_for_image(ocr_text, vision_text, ocr_confidence, vision_confidence, image_id)
            elif ocr_success and not vision_success:
                # 仅OCR成功
                logger.info(f"图像{image_id}使用OCR结果")
                return ocr_text, "ocr_only"
            elif not ocr_success and vision_success:
                # 仅Qwen-VL成功
                logger.info(f"图像{image_id}使用Qwen-VL结果")
                return vision_text, "vision_only"
            else:
                # 两者都失败
                ocr_error = ocr_result.get('error', '未执行') if ocr_result else '未执行'
                vision_error = vision_result.get('error', '未执行') if vision_result else '未执行'
                logger.warning(f"图像{image_id}OCR和Qwen-VL都失败 - OCR: {ocr_error}, Vision: {vision_error}")
                return "", "both_failed"  # 返回空字符串，让上层过滤掉
                
        except Exception as e:
            logger.error(f"图像{image_id}融合决策模块失败: {str(e)}")
            return "", "fusion_error"  # 返回空字符串，让上层过滤掉
    
    def _intelligent_fusion_for_image(self, ocr_text: str, vision_text: str, 
                                    ocr_confidence: float, vision_confidence: float, 
                                    image_id: str) -> Tuple[str, str]:
        """
        智能融合策略 - 针对单个图像
        
        Args:
            ocr_text: OCR识别文本
            vision_text: Qwen-VL识别文本
            ocr_confidence: OCR置信度
            vision_confidence: Qwen-VL置信度
            image_id: 图像ID
            
        Returns:
            (融合文本, 提取方法)
        """
        try:
            # 计算文本质量指标
            ocr_length = len(ocr_text)
            vision_length = len(vision_text)
            
            # 置信度权重
            confidence_weight = 0.4
            length_weight = 0.3
            quality_weight = 0.3
            
            # 计算综合得分
            ocr_score = (ocr_confidence * confidence_weight + 
                        min(ocr_length / 100, 1.0) * length_weight + 
                        self._calculate_text_quality(ocr_text) * quality_weight)
            
            vision_score = (vision_confidence * confidence_weight + 
                           min(vision_length / 100, 1.0) * length_weight + 
                           self._calculate_text_quality(vision_text) * quality_weight)
            
            # 决策策略
            if abs(ocr_score - vision_score) < 0.1:
                # 得分相近：合并结果
                logger.info(f"图像{image_id}OCR和Qwen-VL得分相近，进行合并")
                merged_text = self._merge_texts(ocr_text, vision_text)
                return merged_text, "intelligent_merge"
            elif ocr_score > vision_score:
                # OCR得分更高：以OCR为主，Qwen-VL补充
                logger.info(f"图像{image_id}OCR得分更高，以OCR为主")
                enhanced_text = self._enhance_with_vision(ocr_text, vision_text)
                return enhanced_text, "ocr_enhanced"
            else:
                # Qwen-VL得分更高：以Qwen-VL为主，OCR补充
                logger.info(f"图像{image_id}Qwen-VL得分更高，以Qwen-VL为主")
                enhanced_text = self._enhance_with_ocr(vision_text, ocr_text)
                return enhanced_text, "vision_enhanced"
                
        except Exception as e:
            logger.error(f"图像{image_id}智能融合失败: {str(e)}")
            # 降级到简单合并
            return f"{ocr_text}\n\n[Qwen-VL补充]\n{vision_text}", "simple_merge"
    
    def _calculate_text_quality(self, text: str) -> float:
        """
        计算文本质量得分
        
        Args:
            text: 输入文本
            
        Returns:
            质量得分 (0-1)
        """
        try:
            if not text.strip():
                return 0.0
            
            # 基础指标
            length_score = min(len(text) / 200, 1.0)  # 长度得分
            
            # 字符多样性
            unique_chars = len(set(text))
            diversity_score = min(unique_chars / 50, 1.0)
            
            # 中文字符比例（中文文档质量指标）
            chinese_chars = sum(1 for char in text if '\u4e00' <= char <= '\u9fff')
            chinese_ratio = chinese_chars / len(text) if text else 0
            chinese_score = min(chinese_ratio * 2, 1.0)  # 中文比例得分
            
            # 数字和标点符号（结构化内容指标）
            digit_punct_count = sum(1 for char in text if char.isdigit() or char in '.,;:!?()[]{}')
            structure_score = min(digit_punct_count / 20, 1.0)
            
            # 综合得分
            quality_score = (length_score * 0.3 + 
                           diversity_score * 0.2 + 
                           chinese_score * 0.3 + 
                           structure_score * 0.2)
            
            return min(quality_score, 1.0)
            
        except Exception as e:
            logger.error(f"文本质量计算失败: {str(e)}")
            return 0.5  # 默认中等质量
    
    def _format_user_friendly_text(self, text: str, extraction_method: str) -> str:
        """
        格式化用户友好的文本，过滤技术细节和失败信息
        
        Args:
            text: 原始文本
            extraction_method: 提取方法
            
        Returns:
            格式化后的用户友好文本，如果无意义则返回空字符串
        """
        if not text or not text.strip():
            return ""
        
        # 过滤失败信息和技术细节
        failure_indicators = [
            "处理失败", "未执行", "融合决策失败", "both_failed", "fusion_error",
            "图中没有", "图片中没有", "无可见文字", "无法提取", "没有清晰可见",
            "图中所有文字：", "图片中的文字内容如下：", "这是图片中唯一可见的文字内容",
            "由于图片中没有明显的文字内容", "因此无法提取任何文字"
        ]
        
        # 检查是否包含失败信息
        text_lower = text.lower()
        for indicator in failure_indicators:
            if indicator in text_lower:
                return ""  # 过滤掉失败信息
        
        # 过滤纯技术信息
        if extraction_method in ["both_failed", "fusion_error"]:
            return ""
        
        # 清理文本格式
        cleaned_text = text.strip()
        
        # 移除多余的空行和换行符
        cleaned_text = '\n'.join(line.strip() for line in cleaned_text.split('\n') if line.strip())
        
        # 过滤过短的文本（少于3个字符）
        if len(cleaned_text) < 3:
            return ""
        
        # 过滤纯数字或纯符号
        if cleaned_text.isdigit() or all(c in '0123456789 \n\t' for c in cleaned_text):
            return ""
        
        # 过滤重复内容（如"奖牌"重复很多次）
        words = cleaned_text.split()
        if len(words) > 10:  # 如果单词数较多
            word_counts = {}
            for word in words:
                word_counts[word] = word_counts.get(word, 0) + 1
            
            # 如果某个词出现频率过高，可能是重复内容
            max_freq = max(word_counts.values()) if word_counts else 0
            if max_freq > len(words) * 0.3:  # 如果某个词出现超过30%
                return ""
        
        return cleaned_text
    
    def _merge_texts(self, ocr_text: str, vision_text: str) -> str:
        """
        合并两个文本结果
        
        Args:
            ocr_text: OCR文本
            vision_text: Qwen-VL文本
            
        Returns:
            合并后的文本
        """
        try:
            # 简单的合并策略：去重并保持顺序
            lines_ocr = [line.strip() for line in ocr_text.split('\n') if line.strip()]
            lines_vision = [line.strip() for line in vision_text.split('\n') if line.strip()]
            
            # 合并并去重
            merged_lines = []
            seen = set()
            
            for line in lines_ocr + lines_vision:
                if line not in seen:
                    merged_lines.append(line)
                    seen.add(line)
            
            return '\n'.join(merged_lines)
            
        except Exception as e:
            logger.error(f"文本合并失败: {str(e)}")
            return f"{ocr_text}\n\n{vision_text}"
    
    def _enhance_with_vision(self, ocr_text: str, vision_text: str) -> str:
        """
        以OCR为主，用Qwen-VL增强
        
        Args:
            ocr_text: 主要文本（OCR）
            vision_text: 增强文本（Qwen-VL）
            
        Returns:
            增强后的文本
        """
        try:
            # 如果OCR文本已经很完整，直接返回
            if len(ocr_text) > len(vision_text) * 1.5:
                return ocr_text
            
            # 否则进行补充
            return f"{ocr_text}\n\n[Qwen-VL补充信息]\n{vision_text}"
            
        except Exception as e:
            logger.error(f"OCR增强失败: {str(e)}")
            return ocr_text
    
    def _enhance_with_ocr(self, vision_text: str, ocr_text: str) -> str:
        """
        以Qwen-VL为主，用OCR增强
        
        Args:
            vision_text: 主要文本（Qwen-VL）
            ocr_text: 增强文本（OCR）
            
        Returns:
            增强后的文本
        """
        try:
            # 如果Qwen-VL文本已经很完整，直接返回
            if len(vision_text) > len(ocr_text) * 1.5:
                return vision_text
            
            # 否则进行补充
            return f"{vision_text}\n\n[OCR补充信息]\n{ocr_text}"
            
        except Exception as e:
            logger.error(f"Qwen-VL增强失败: {str(e)}")
            return vision_text
    
    async def extract_text_from_doc_async(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """
        异步从DOC/DOCX文件提取文字
        
        Args:
            file_path: 文件路径
            **kwargs: 其他参数
            
        Returns:
            提取结果
        """
        try:
            # 提取参数
            doc_type_str = kwargs.get('doc_type', 'docx')
            use_ocr = kwargs.get('use_ocr', False)
            ocr_engine = kwargs.get('ocr_engine', 'baidu')
            use_vision = kwargs.get('use_vision', False)
            vision_model = kwargs.get('vision_model', 'qwen-vl-plus')
            
            logger.info(f"DOC处理参数: doc_type={doc_type_str}, use_ocr={use_ocr}, ocr_engine={ocr_engine}, use_vision={use_vision}, vision_model={vision_model}")
            
            # 转换文档类型
            try:
                doc_type = DocType(doc_type_str.lower())
            except ValueError:
                logger.warning(f"不支持的文档类型: {doc_type_str}，使用默认类型: docx")
                doc_type = DocType.DOCX
            
            # 转换OCR引擎类型
            try:
                ocr_engine_enum = OCREngineType(ocr_engine.lower())
            except ValueError:
                logger.warning(f"不支持的OCR引擎类型: {ocr_engine}，使用默认引擎: baidu")
                ocr_engine_enum = OCREngineType.BAIDU
            
            # 提取文字
            result = self.extract_text_from_doc(
                file_path, 
                doc_type=doc_type,
                use_ocr=use_ocr, 
                ocr_engine=ocr_engine_enum,
                use_vision=use_vision,
                vision_model=vision_model,
            )
            
            # 转换为字典格式
            return result.dict()
            
        except Exception as e:
            logger.error(f"DOC文字提取失败: {str(e)}")
            raise


# 创建全局实例
doc_extractor = DocTextExtractor()
